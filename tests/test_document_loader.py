from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
import pytest

from src.document_loader import DocumentLoadError, load_directory, load_docx


def make_docx(path: Path) -> None:
    document = Document()
    document.add_heading("FastAPI 入门", level=1)
    document.add_paragraph("  使用   类型注解。  ")
    document.add_paragraph("")
    document.save(path)


def corrupt_document_xml(path: Path) -> None:
    with ZipFile(path, "r") as archive:
        files = {
            name: archive.read(name)
            for name in archive.namelist()
        }
    files["word/document.xml"] = b"<w:document"
    with ZipFile(path, "w", compression=ZIP_DEFLATED) as archive:
        for name, content in files.items():
            archive.writestr(name, content)


def test_load_docx_normalizes_text_and_keeps_positions(tmp_path: Path) -> None:
    path = tmp_path / "guide.docx"
    make_docx(path)

    paragraphs = load_docx(path)

    assert [item.text for item in paragraphs] == [
        "FastAPI 入门",
        "使用 类型注解。",
    ]
    assert [item.position for item in paragraphs] == [1, 2]
    assert paragraphs[0].is_heading is True
    assert paragraphs[0].source == "guide.docx"


def test_load_docx_recognizes_localized_heading_style(tmp_path: Path) -> None:
    path = tmp_path / "localized.docx"
    document = Document()
    style = document.styles.add_style("标题自定义", WD_STYLE_TYPE.PARAGRAPH)
    paragraph = document.add_paragraph("中文标题")
    paragraph.style = style
    document.save(path)

    paragraphs = load_docx(path)

    assert paragraphs[0].is_heading is True


def test_load_docx_rejects_empty_document(tmp_path: Path) -> None:
    path = tmp_path / "empty.docx"
    Document().save(path)

    with pytest.raises(DocumentLoadError, match="empty"):
        load_docx(path)


def test_load_docx_rejects_legacy_doc_format(tmp_path: Path) -> None:
    path = tmp_path / "legacy.doc"
    path.write_bytes(b"legacy Word content")

    with pytest.raises(DocumentLoadError, match="unsupported file type: .doc"):
        load_docx(path)


def test_load_directory_uses_sorted_docx_files(tmp_path: Path) -> None:
    make_docx(tmp_path / "b.docx")
    make_docx(tmp_path / "a.docx")
    (tmp_path / "ignored.txt").write_text("not a Word document", encoding="utf-8")

    documents, errors = load_directory(tmp_path)

    assert list(documents) == ["a.docx", "b.docx"]
    assert errors == {}


def test_load_directory_records_corrupt_file_and_keeps_valid_files(
    tmp_path: Path,
) -> None:
    make_docx(tmp_path / "valid.docx")
    (tmp_path / "broken.docx").write_bytes(b"not a zip package")

    documents, errors = load_directory(tmp_path)

    assert list(documents) == ["valid.docx"]
    assert "broken.docx" in errors
    assert "cannot read" in errors["broken.docx"]


def test_load_directory_records_malformed_word_xml(
    tmp_path: Path,
) -> None:
    make_docx(tmp_path / "valid.docx")
    malformed_path = tmp_path / "malformed.docx"
    make_docx(malformed_path)
    corrupt_document_xml(malformed_path)

    documents, errors = load_directory(tmp_path)

    assert list(documents) == ["valid.docx"]
    assert "malformed.docx" in errors
    assert "cannot read" in errors["malformed.docx"]
